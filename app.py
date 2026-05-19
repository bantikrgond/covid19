import os
import sqlite3
import pandas as pd
from io import BytesIO
from functools import wraps
from flask import Flask, render_template, request, redirect, session, jsonify, send_from_directory, Response
from werkzeug.security import generate_password_hash, check_password_hash

app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'covid_prediction_dashboard_secure_key_2026')

# On Vercel, the file system is read-only. Store SQLite database in /tmp.
if os.environ.get('VERCEL') or os.environ.get('NOW_REGION'):
    DB_PATH = '/tmp/predictions.db'
else:
    DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'predictions.db')

def get_db_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    """Initializes the database, creating the predictions and admins tables, and sets up a default admin."""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # 1. Create Predictions Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS predictions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            age INTEGER NOT NULL,
            gender TEXT NOT NULL,
            fever TEXT NOT NULL,
            cough TEXT NOT NULL,
            fatigue TEXT NOT NULL,
            chest_pain TEXT NOT NULL,
            breathing_difficulty TEXT NOT NULL,
            diabetes TEXT NOT NULL,
            risk_level TEXT NOT NULL,
            confidence TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # 2. Create Admins Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS admins (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL
        )
    ''')
    
    # 3. Insert default admin if table is empty or does not have bantikr11
    cursor.execute('SELECT COUNT(*) FROM admins WHERE username = ?', ('bantikr11',))
    if cursor.fetchone()[0] == 0:
        # Clear existing admins to reset to user's requested config
        cursor.execute('DELETE FROM admins')
        default_username = 'bantikr11'
        default_password = 'Banti123'
        hashed_password = generate_password_hash(default_password)
        cursor.execute(
            'INSERT INTO admins (username, password_hash) VALUES (?, ?)',
            (default_username, hashed_password)
        )
        print("=" * 60)
        print("DATABASE INITIALIZED: Default Administrator Created!")
        print("Username: bantikr11")
        print("Password: Banti123")
        print("WARNING: Please change this password in a production environment.")
        print("=" * 60)
        
    # 4. Create Customers Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS customers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            name TEXT NOT NULL
        )
    ''')

    # 5. Create Customer Health Data Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS customer_health_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            customer_id INTEGER NOT NULL,
            metric_name TEXT NOT NULL,
            previous_value REAL NOT NULL,
            current_value REAL NOT NULL,
            unit TEXT NOT NULL,
            FOREIGN KEY(customer_id) REFERENCES customers(id)
        )
    ''')

    # Insert default customer if table is empty
    cursor.execute('SELECT COUNT(*) FROM customers WHERE email = ?', ('customer@covidai.com',))
    if cursor.fetchone()[0] == 0:
        hashed_password = generate_password_hash('Customer123')
        cursor.execute(
            'INSERT INTO customers (email, password_hash, name) VALUES (?, ?, ?)',
            ('customer@covidai.com', hashed_password, 'Alice Smith')
        )
        customer_id = cursor.lastrowid
        
        # Insert health data
        health_metrics = [
            ('Body Temperature', 101.5, 98.6, 'F'),
            ('Symptom Severity', 8.0, 2.0, 'score'),
            ('Oxygen Saturation', 91.0, 98.0, 'percent'),
            ('Heart Rate', 95.0, 72.0, 'bpm')
        ]
        for metric, prev_val, curr_val, unit in health_metrics:
            cursor.execute(
                'INSERT INTO customer_health_data (customer_id, metric_name, previous_value, current_value, unit) VALUES (?, ?, ?, ?, ?)',
                (customer_id, metric, prev_val, curr_val, unit)
            )
        print("=" * 60)
        print("DATABASE INITIALIZED: Default Customer Created!")
        print("Email: customer@covidai.com")
        print("Password: Customer123")
        print("=" * 60)

    conn.commit()
    conn.close()

# Initialize the database on startup
init_db()

@app.after_request
def add_cors_headers(response):
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type,Authorization'
    response.headers['Access-Control-Allow-Methods'] = 'GET,PUT,POST,DELETE,OPTIONS'
    return response

# ==========================================
# AUTHENTICATION DECORATOR
# ==========================================
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'admin_logged_in' not in session:
            # Check if AJAX/API request or standard page request
            if request.path.startswith('/admin/api/') or request.path.startswith('/admin/export/'):
                return jsonify({"error": "Unauthorized. Login required."}), 401
            return redirect('/admin/login')
        return f(*args, **kwargs)
    return decorated_function

# ==========================================
# FRONTEND / USER ROUTES
# ==========================================

@app.route('/')
def index():
    """Serves the main landing page."""
    return send_from_directory('.', 'index.html')

@app.route('/script.js')
def serve_script():
    """Serves the main website javascript file."""
    return send_from_directory('.', 'script.js')

@app.route('/style.css')
def serve_style():
    """Serves the main website stylesheet."""
    return send_from_directory('.', 'style.css')

@app.route('/api/save_prediction', methods=['POST', 'OPTIONS'])
def save_prediction():
    """Endpoint for background prediction data collection."""
    if request.method == 'OPTIONS':
        return Response(status=200)
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "message": "No JSON payload received"}), 400
        
        # Extract fields
        name = data.get('name', 'Anonymous Patient').strip() or 'Anonymous Patient'
        age = int(data.get('age', 0))
        gender = data.get('gender', 'Unknown')
        fever = data.get('fever', 'No')
        cough = data.get('cough', 'No')
        fatigue = data.get('fatigue', 'No')
        chest_pain = data.get('chest', 'No')  # JS key: chest
        breathing_difficulty = data.get('breathing', 'No')  # JS key: breathing
        diabetes = data.get('diabetes', 'No')
        risk_level = data.get('risk', 'LOW RISK')
        confidence = data.get('confidence', '0.0%')
        
        # Save to database
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO predictions (
                name, age, gender, fever, cough, fatigue, chest_pain, breathing_difficulty, diabetes, risk_level, confidence
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (name, age, gender, fever, cough, fatigue, chest_pain, breathing_difficulty, diabetes, risk_level, confidence))
        
        conn.commit()
        conn.close()
        
        return jsonify({"success": True, "message": "Prediction record stored successfully"}), 200
        
    except Exception as e:
        print(f"Error saving prediction: {e}")
        return jsonify({"success": False, "message": "An error occurred while saving prediction data"}), 500

@app.route('/api/stats')
def public_stats():
    """Public endpoint to return database statistics for the homepage."""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT COUNT(*) FROM predictions')
        count = cursor.fetchone()[0]
        conn.close()
        return jsonify({"success": True, "active_users": count}), 200
    except Exception as e:
        print(f"Error fetching stats: {e}")
        return jsonify({"success": False, "message": str(e)}), 500

# ==========================================
# ADMIN AUTH ROUTES
# ==========================================

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    """Handles admin authentication login screen and logic."""
    if 'admin_logged_in' in session:
        return redirect('/admin')
        
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()
        
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM admins WHERE username = ?', (username,))
        admin = cursor.fetchone()
        conn.close()
        
        if admin and check_password_hash(admin['password_hash'], password):
            session['admin_logged_in'] = True
            session['admin_user'] = username
            return redirect('/admin')
        else:
            return render_template('login.html', error="Invalid username or password credentials.")
            
    return render_template('login.html', error=None)

@app.route('/admin/logout')
def admin_logout():
    """Logs the admin user out and destroys session context."""
    session.pop('admin_logged_in', None)
    session.pop('admin_user', None)
    return redirect('/admin/login')

# ==========================================
# CUSTOMER AUTH DECORATOR & ROUTES
# ==========================================
def customer_login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'customer_logged_in' not in session:
            return redirect('/customer/login')
        return f(*args, **kwargs)
    return decorated_function

@app.route('/customer/login', methods=['GET', 'POST'])
def customer_login():
    """Handles customer authentication login screen and logic."""
    if 'customer_logged_in' in session:
        return redirect('/customer/dashboard')
        
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '').strip()
        
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM customers WHERE email = ?', (email,))
        customer = cursor.fetchone()
        conn.close()
        
        if customer and check_password_hash(customer['password_hash'], password):
            session['customer_logged_in'] = True
            session['customer_id'] = customer['id']
            session['customer_email'] = customer['email']
            session['customer_name'] = customer['name']
            return redirect('/customer/dashboard')
        else:
            return render_template('customer_login.html', error="Invalid email or password.")
            
    return render_template('customer_login.html', error=None)

@app.route('/customer/logout')
def customer_logout():
    """Logs the customer out and destroys session context."""
    session.pop('customer_logged_in', None)
    session.pop('customer_id', None)
    session.pop('customer_email', None)
    session.pop('customer_name', None)
    return redirect('/customer/login')

@app.route('/customer/dashboard')
@customer_login_required
def customer_dashboard():
    """Renders the customer health dashboard."""
    customer_id = session['customer_id']
    
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM customer_health_data WHERE customer_id = ?', (customer_id,))
    metrics = cursor.fetchall()
    conn.close()
    
    return render_template('customer_dashboard.html', metrics=metrics)

@app.route('/customer/chart.png')
@customer_login_required
def customer_chart_png():
    """Generates the comparison chart dynamically using matplotlib."""
    customer_id = session['customer_id']
    
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('SELECT metric_name, previous_value, current_value, unit FROM customer_health_data WHERE customer_id = ?', (customer_id,))
    rows = cursor.fetchall()
    conn.close()
    
    # Use standard labels for before/after comparison
    metric_names = [f"{row['metric_name']} ({row['unit']})" for row in rows]
    previous = [row['previous_value'] for row in rows]
    current = [row['current_value'] for row in rows]
        
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np
    
    x = np.arange(len(metric_names))
    width = 0.35
    
    fig, ax = plt.subplots(figsize=(8, 4.5))
    rects1 = ax.bar(x - width/2, previous, width, label='Previous (Before)', color='#ef4444')
    rects2 = ax.bar(x + width/2, current, width, label='Current (After)', color='#10b981')
    
    ax.set_ylabel('Health Metrics')
    ax.set_title('Customer Health Comparison (Before vs After)')
    ax.set_xticks(x)
    ax.set_xticklabels(metric_names)
    ax.legend(loc='upper right')
    
    # Label height adjustments on top of bars
    for rect in rects1:
        height = rect.get_height()
        ax.annotate(f'{height}',
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # 3 points vertical offset
                    textcoords="offset points",
                    ha='center', va='bottom')
                    
    for rect in rects2:
        height = rect.get_height()
        ax.annotate(f'{height}',
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # 3 points vertical offset
                    textcoords="offset points",
                    ha='center', va='bottom')
                    
    fig.tight_layout()
    
    output = BytesIO()
    plt.savefig(output, format='png')
    plt.close(fig)
    output.seek(0)
    
    return Response(output.read(), mimetype='image/png')

# ==========================================
# ADMIN DASHBOARD ROUTES
# ==========================================

@app.route('/admin')
@login_required
def admin_dashboard():
    """Renders the admin analytics and management dashboard."""
    return render_template('admin.html')

@app.route('/admin/api/predictions')
@login_required
def get_predictions():
    """AJAX endpoint to retrieve all stored records."""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        # Sort newest records first
        cursor.execute('SELECT * FROM predictions ORDER BY created_at DESC')
        rows = cursor.fetchall()
        conn.close()
        
        records = []
        for row in rows:
            records.append({
                "id": row["id"],
                "name": row["name"],
                "age": row["age"],
                "gender": row["gender"],
                "fever": row["fever"],
                "cough": row["cough"],
                "fatigue": row["fatigue"],
                "chest_pain": row["chest_pain"],
                "breathing_difficulty": row["breathing_difficulty"],
                "diabetes": row["diabetes"],
                "risk_level": row["risk_level"],
                "confidence": row["confidence"],
                "created_at": row["created_at"]
            })
            
        return jsonify({"success": True, "records": records}), 200
    except Exception as e:
        print(f"Error fetching predictions: {e}")
        return jsonify({"success": False, "message": str(e)}), 500

@app.route('/admin/api/predictions/<int:record_id>/delete', methods=['POST'])
@login_required
def delete_prediction(record_id):
    """Deletes a clinical prediction entry from the database."""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('DELETE FROM predictions WHERE id = ?', (record_id,))
        conn.commit()
        conn.close()
        return jsonify({"success": True, "message": "Record deleted successfully"}), 200
    except Exception as e:
        print(f"Error deleting record: {e}")
        return jsonify({"success": False, "message": str(e)}), 500

@app.route('/admin/export/<string:format_type>')
@login_required
def export_data(format_type):
    """Exports clinical records to CSV, Excel, PDF, or Word file formats."""
    try:
        conn = get_db_connection()
        # Read the sqlite table into a pandas DataFrame
        df = pd.read_sql_query('SELECT * FROM predictions ORDER BY created_at DESC', conn)
        conn.close()
        
        # Rename columns for a clean presentation
        df.columns = [
            'Record ID', 'Patient Name', 'Age', 'Gender', 'Fever', 'Cough', 
            'Fatigue', 'Chest Pain', 'Difficulty Breathing', 'Diabetes', 
            'Predicted Risk Level', 'Confidence Score', 'Submission Date'
        ]
        
        if format_type == 'csv':
            # Export CSV directly
            csv_data = df.to_csv(index=False)
            return Response(
                csv_data,
                mimetype="text/csv",
                headers={"Content-disposition": "attachment; filename=covid_predictions_export.csv"}
            )
            
        elif format_type == 'xlsx':
            # Export Excel using pandas and openpyxl
            output = BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Clinical Records')
            output.seek(0)
            
            return Response(
                output.read(),
                mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-disposition": "attachment; filename=covid_predictions_export.xlsx"}
            )
            
        elif format_type == 'pdf':
            # Export PDF using ReportLab
            from reportlab.lib.pagesizes import letter, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors
            
            output = BytesIO()
            doc = SimpleDocTemplate(
                output, 
                pagesize=landscape(letter),
                rightMargin=20, leftMargin=20, topMargin=20, bottomMargin=20
            )
            
            elements = []
            styles = getSampleStyleSheet()
            
            # Title style
            title_style = ParagraphStyle(
                'TitleStyle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#0f172a'),
                spaceAfter=15
            )
            elements.append(Paragraph("COVID-19 Clinical Prediction Records Report", title_style))
            elements.append(Spacer(1, 10))
            
            headers = [
                'Submission Date', 'Patient Name', 'Age', 'Gender', 
                'Fever', 'Cough', 'Fatigue', 'Chest Pain', 'Breathing', 'Diabetes', 
                'Risk Level', 'Confidence'
            ]
            
            data_rows = [headers]
            
            cell_text_style = ParagraphStyle(
                'CellTextStyle',
                parent=styles['Normal'],
                fontSize=8,
                leading=10
            )
            header_text_style = ParagraphStyle(
                'HeaderTextStyle',
                parent=styles['Normal'],
                fontSize=8,
                leading=10,
                textColor=colors.white,
                fontName='Helvetica-Bold'
            )
            
            for idx, r in df.iterrows():
                timestamp_str = str(r['Submission Date'])
                if len(timestamp_str) > 19:
                    timestamp_str = timestamp_str[:16]
                
                risk_lvl = str(r['Predicted Risk Level'])
                
                row = [
                    Paragraph(timestamp_str, cell_text_style),
                    Paragraph(str(r['Patient Name']), cell_text_style),
                    Paragraph(str(r['Age']), cell_text_style),
                    Paragraph(str(r['Gender']), cell_text_style),
                    Paragraph(str(r['Fever']), cell_text_style),
                    Paragraph(str(r['Cough']), cell_text_style),
                    Paragraph(str(r['Fatigue']), cell_text_style),
                    Paragraph(str(r['Chest Pain']), cell_text_style),
                    Paragraph(str(r['Difficulty Breathing']), cell_text_style),
                    Paragraph(str(r['Diabetes']), cell_text_style),
                    Paragraph(risk_lvl, cell_text_style),
                    Paragraph(str(r['Confidence Score']), cell_text_style)
                ]
                data_rows.append(row)
                
            col_widths = [100, 100, 30, 45, 35, 35, 35, 45, 45, 45, 80, 57]
            
            t = Table(data_rows, colWidths=col_widths, repeatRows=1)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e293b')),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                ('TOPPADDING', (0,0), (-1,-1), 6),
                ('LEFTPADDING', (0,0), (-1,-1), 4),
                ('RIGHTPADDING', (0,0), (-1,-1), 4),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
            ]))
            
            for i in range(1, len(data_rows)):
                risk_cell = df.iloc[i-1]['Predicted Risk Level']
                if 'HIGH' in risk_cell.upper():
                    t.setStyle(TableStyle([('BACKGROUND', (10, i), (10, i), colors.HexColor('#fee2e2')),
                                           ('TEXTCOLOR', (10, i), (10, i), colors.HexColor('#991b1b'))]))
                elif 'MEDIUM' in risk_cell.upper():
                    t.setStyle(TableStyle([('BACKGROUND', (10, i), (10, i), colors.HexColor('#fef3c7')),
                                           ('TEXTCOLOR', (10, i), (10, i), colors.HexColor('#92400e'))]))
                else:
                    t.setStyle(TableStyle([('BACKGROUND', (10, i), (10, i), colors.HexColor('#dcfce7')),
                                           ('TEXTCOLOR', (10, i), (10, i), colors.HexColor('#166534'))]))
            
            for c_idx in range(len(headers)):
                data_rows[0][c_idx] = Paragraph(headers[c_idx], header_text_style)
                
            elements.append(t)
            doc.build(elements)
            output.seek(0)
            
            return Response(
                output.read(),
                mimetype="application/pdf",
                headers={"Content-disposition": "attachment; filename=covid_predictions_export.pdf"}
            )
            
        elif format_type == 'docx':
            # Export Word using python-docx
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Landscape layout
            section = doc.sections[0]
            section.orientation = 1  # Landscape
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            
            title = doc.add_heading("COVID-19 Clinical Prediction Records Report", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph("This document contains the clinical risk predictions generated by the COVID AI model.")
            
            headers = [
                'Submission Date', 'Patient Name', 'Age', 'Gender', 
                'Fever', 'Cough', 'Fatigue', 'Chest Pain', 'Breathing', 'Diabetes', 
                'Risk Level', 'Confidence'
            ]
            
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Shading Accent 1'
            
            hdr_cells = table.rows[0].cells
            for idx, text in enumerate(headers):
                hdr_cells[idx].text = text
                for p in hdr_cells[idx].paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)
                        
            for idx, r in df.iterrows():
                row_cells = table.add_row().cells
                
                timestamp_str = str(r['Submission Date'])
                if len(timestamp_str) > 19:
                    timestamp_str = timestamp_str[:16]
                    
                row_cells[0].text = timestamp_str
                row_cells[1].text = str(r['Patient Name'])
                row_cells[2].text = str(r['Age'])
                row_cells[3].text = str(r['Gender'])
                row_cells[4].text = str(r['Fever'])
                row_cells[5].text = str(r['Cough'])
                row_cells[6].text = str(r['Fatigue'])
                row_cells[7].text = str(r['Chest Pain'])
                row_cells[8].text = str(r['Difficulty Breathing'])
                row_cells[9].text = str(r['Diabetes'])
                row_cells[10].text = str(r['Predicted Risk Level'])
                row_cells[11].text = str(r['Confidence Score'])
                
                for cell in row_cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(8.5)
                            
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            return Response(
                output.read(),
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-disposition": "attachment; filename=covid_predictions_export.docx"}
            )
            
        else:
            return jsonify({"success": False, "message": "Invalid export format specified. Must be csv, xlsx, pdf, or docx."}), 400
            
    except Exception as e:
        print(f"Error exporting data: {e}")
        return jsonify({"success": False, "message": str(e)}), 500

if __name__ == '__main__':
    # Run server locally on port 5000
    app.run(host='127.0.0.1', port=5000, debug=True)
